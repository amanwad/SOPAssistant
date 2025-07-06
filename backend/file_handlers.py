import os
from pathlib import Path
from typing import Dict, Any, Optional
import logging

# Import libraries for different file types
import pymupdf
import docx
import pandas as pd
from PIL import Image
import pytesseract

from embedding import split_and_upsert

class FileHandler:
    """Handler class for processing different file types and extracting text content."""
    
    def __init__(self):
        self.supported_extensions = {
            '.txt': self._handle_text_file,
            '.md': self._handle_text_file,
            '.pdf': self._handle_pdf_file,
            '.docx': self._handle_docx_file,
            '.doc': self._handle_doc_file,
            '.xlsx': self._handle_excel_file,
            '.xls': self._handle_excel_file,
            '.csv': self._handle_csv_file,
            '.jpg': self._handle_image_file,
            '.jpeg': self._handle_image_file,
            '.png': self._handle_image_file,
            '.gif': self._handle_image_file,
            '.bmp': self._handle_image_file,
            '.tiff': self._handle_image_file,
        }
        uploads_dir = Path(__file__).parent / "uploads"
        if uploads_dir.exists() and uploads_dir.is_dir():
            self.doc_num = len([f for f in uploads_dir.iterdir() if f.is_file()])
        else:
            self.doc_num = 0
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        Process a file and extract its text content based on file type.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        print(file_path)
        file_path = Path(file_path)

        
        
        if not file_path.exists():
            return {
                'success': False,
                'error': f'File not found: {file_path}',
                'text': '',
                'metadata': {}
            }
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_extensions:
            return {
                'success': False,
                'error': f'Unsupported file type: {file_extension}',
                'text': '',
                'metadata': {
                    'file_type': file_extension,
                    'file_size': file_path.stat().st_size,
                    'file_name': file_path.name
                }
            }
        
        try:
            handler_func = self.supported_extensions[file_extension]
            result = handler_func(file_path)
            print("process file handler", result)
            result['metadata']['file_type'] = file_extension
            result['metadata']['file_size'] = file_path.stat().st_size
            result['metadata']['file_name'] = file_path.name
            self.doc_num += 1
            return result
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing file: {str(e)}',
                'text': '',
                'metadata': {
                    'file_type': file_extension,
                    'file_size': file_path.stat().st_size,
                    'file_name': file_path.name
                }
            }
    
    def _handle_text_file(self, file_path: Path) -> Dict[str, Any]:
        """Handle plain text files (.txt, .md, etc.)"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return {
                'success': True,
                'text': text,
                'metadata': {
                    'encoding': 'utf-8',
                    'line_count': len(text.splitlines()),
                    'word_count': len(text.split())
                }
            }
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                return {
                    'success': True,
                    'text': text,
                    'metadata': {
                        'encoding': 'latin-1',
                        'line_count': len(text.splitlines()),
                        'word_count': len(text.split())
                    }
                }
            except Exception as e:
                return {
                    'success': False,
                    'error': f'Failed to read text file: {str(e)}',
                    'text': ''
                }
    
    def _handle_pdf_file(self, file_path: Path) -> Dict[str, Any]:
        """Handle PDF files using PyPDF2"""
        try:
            text = ""
            file = pymupdf.open(file_path)
            for i, page in enumerate(file):
                print("iterating")
                metadata = {
                    "page_number": i,
                    "file_name": file_path.name,
                    "file_path": str(file_path),
                    "file_size": file_path.stat().st_size,
                    "file_type": file_path.suffix,
                    "doc_num": self.doc_num
                }
                print(metadata)
                split_and_upsert(page.get_text(), metadata)
                print("upsert")
                text += page.get_text() + "\n"
            
            print("153")
            page_count = file.page_count if hasattr(file, "page_count") else len(file)
            result = {
                'success': True,
                'text': text.strip(),
                'metadata': {
                    'page_count': page_count,
                    'line_count': len(text.splitlines()),
                    'word_count': len(text.split())
                }
            }
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to extract text from PDF: {str(e)}',
                'text': ''
            }
    
    def _handle_docx_file(self, file_path: Path) -> Dict[str, Any]:
        """Handle DOCX files using python-docx"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for i, paragraph in enumerate(doc.paragraphs):
                metadata = {
                    "paragraph_number": i,
                    "file_name": file_path.name,
                    "file_path": file_path,
                    "file_size": file_path.stat().st_size,
                    "file_type": file_path.suffix,
                    "doc_num": self.doc_num
                }
                split_and_upsert(paragraph.text, metadata)
                text += paragraph.text + "\n"
            
            return {
                'success': True,
                'text': text.strip(),
                'metadata': {
                    'paragraph_count': len(doc.paragraphs),
                    'line_count': len(text.splitlines()),
                    'word_count': len(text.split())
                }
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to extract text from DOCX: {str(e)}',
                'text': ''
            }
    
    def _handle_doc_file(self, file_path: Path) -> Dict[str, Any]:
        """Handle DOC files (legacy Word format)"""
        # Note: This is a placeholder. Processing .doc files requires additional libraries
        # like antiword or textract. For now, we'll return an error.
        return {
            'success': False,
            'error': 'DOC file processing not implemented. Please convert to DOCX format.',
            'text': ''
        }
    
    def _handle_excel_file(self, file_path: Path) -> Dict[str, Any]:
        """Handle Excel files (.xlsx, .xls) using pandas"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text = ""
            sheet_names = excel_file.sheet_names
            
            for sheet_name in sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text += f"Sheet: {sheet_name}\n"
                text += df.to_string(index=False) + "\n\n"
            
            return {
                'success': True,
                'text': text.strip(),
                'metadata': {
                    'sheet_count': len(sheet_names),
                    'sheet_names': sheet_names,
                    'line_count': len(text.splitlines()),
                    'word_count': len(text.split())
                }
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to extract text from Excel file: {str(e)}',
                'text': ''
            }
    
    def _handle_csv_file(self, file_path: Path) -> Dict[str, Any]:
        """Handle CSV files using pandas"""
        try:
            df = pd.read_csv(file_path)
            text = df.to_string(index=False)
            
            return {
                'success': True,
                'text': text,
                'metadata': {
                    'row_count': len(df),
                    'column_count': len(df.columns),
                    'columns': df.columns.tolist(),
                    'line_count': len(text.splitlines()),
                    'word_count': len(text.split())
                }
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to extract text from CSV file: {str(e)}',
                'text': ''
            }
    
    def _handle_image_file(self, file_path: Path) -> Dict[str, Any]:
        """Handle image files using OCR (Optical Character Recognition)"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            return {
                'success': True,
                'text': text.strip(),
                'metadata': {
                    'image_size': image.size,
                    'image_mode': image.mode,
                    'line_count': len(text.splitlines()),
                    'word_count': len(text.split())
                }
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to extract text from image using OCR: {str(e)}',
                'text': ''
            }
    
    

    def get_supported_formats(self) -> list:
        """Get list of supported file formats"""
        return list(self.supported_extensions.keys())
    
    def is_supported(self, file_extension: str) -> bool:
        """Check if a file extension is supported"""
        return file_extension.lower() in self.supported_extensions 